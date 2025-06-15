import sqlite3
import tkinter as tk
from tkinter import messagebox, filedialog
from customtkinter import CTk, CTkButton, CTkFrame, CTkLabel, CTkEntry, CTkScrollableFrame
from PIL import ImageTk, Image, ExifTags
import os
import shutil
import uuid
import sys
import datetime
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import defaultdict

def resource_path(relative_path):
    """Get absolute path to resource, works for PyInstaller and development"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

DB_FILENAME = resource_path("database/vswa_db.db")

def set_cell_width(cell, width_dxa):
    """
    Set the width of a cell.
    width_dxa: width in dxa units (1 inch = 1440 dxa)
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

class HomepageWindow(CTk):
    def __init__(self, user_id, login_app):
        super().__init__()
        self.title("Homepage - Contractor")
        self.login_app = login_app
        self.current_user_id = user_id
        self.db_filename = DB_FILENAME
        self.add_image_window = None  # Reference to the image upload window
        self.protocol("WM_DELETE_WINDOW", self.on_close)
        self.setup_window()
        self.create_widgets()
        
    def get_last_upload_date(self, item_number, item_name):
        try:
            conn = sqlite3.connect(self.db_filename)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT MAX(upload_date)
                FROM completed_construction_images
                WHERE item_number = ? AND item_name = ?
            """, (item_number, item_name))
            result = cursor.fetchone()
            if result and result[0]:
                iso_date = result[0]
                formatted_date = datetime.datetime.strptime(iso_date, "%Y-%m-%d").strftime("%B %d, %Y")
                return formatted_date
            return "N/A"
        except Exception as e:
            print(f"Error fetching last upload date: {e}")
            return "N/A"
        finally:
            if conn:
                conn.close()
    
    def refresh_table(self):
        self.table_container.destroy()
        self.table_container = tk.Frame(self.extra_frame, bd=2, relief="solid")
        self.table_container.pack(padx=10, pady=10, fill="both", expand=True)
        self.create_custom_table()
        
    # ─── MOUSE WHEEL SCROLLING METHODS ───────────────────────────────
    def _bind_mousewheel(self, event):
        self.table_canvas.bind_all("<MouseWheel>", self._on_mousewheel)
    
    def _unbind_mousewheel(self, event):
        self.table_canvas.unbind_all("<MouseWheel>")
    
    def _on_mousewheel(self, event):
        self.table_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def setup_window(self):
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        window_width = int(screen_width * 0.75)
        window_height = int(screen_height * 0.75)
        position_x = (screen_width - window_width) // 2
        position_y = (screen_height - window_height) // 2

        self.geometry(f"{window_width}x{window_height}+{position_x}+{position_y}")
        self.resizable(False, False)

    def create_widgets(self):
        project_details = self.fetch_project_details()
        project_name = project_details.get("project_name", "N/A")
        location = project_details.get("location", "N/A")
        project_id = project_details.get("project_id", "N/A")
        contractor_name = project_details.get("contractor_name", "N/A")
        project_type = project_details.get("project_type", "N/A")

        self.logout_button = CTkButton(self, text="Logout", font=("Roboto", 12), command=self.logout)
        self.logout_button.pack(pady=5, anchor="ne", padx=10)

        self.info_frame = CTkFrame(self, corner_radius=10, fg_color="#11151f")
        self.info_frame.pack(pady=10, padx=10, fill="x")

        labels_info = [
            f"NAME OF PROJECT: {project_name}",
            f"LOCATION: {location}",
            f"PROJECT ID: {project_id}",
            f"CONTRACTOR: {contractor_name}",
            f"PROJECT TYPE: {project_type}"
        ]
        for text in labels_info:
            CTkLabel(self.info_frame, text=text, font=("Roboto", 12, "bold"), text_color="white", anchor="w").pack(padx=10, pady=2, anchor="w")

        self.extra_frame = CTkFrame(self, corner_radius=10, fg_color="#11151f")
        self.extra_frame.pack(pady=10, padx=10, fill="both", expand=True)

        self.extra_label = CTkLabel(self.extra_frame, text="List of Construction Items", font=("Roboto", 16, "bold"), text_color="white")
        self.extra_label.pack(pady=10)

        self.table_container = tk.Frame(self.extra_frame, bd=2, relief="solid")
        self.table_container.pack(padx=10, pady=10, fill="both", expand=True)

        self.create_custom_table()

    def create_custom_table(self):
        self.table_canvas = tk.Canvas(self.table_container, bg="#1C1C1C", highlightthickness=0)
        self.scrollbar = tk.Scrollbar(self.table_container, orient="vertical", command=self.table_canvas.yview)
        self.scrollable_frame = CTkFrame(self.table_canvas, fg_color="#1C1C1C")

        self.scrollable_frame.bind("<Configure>", lambda e: self.table_canvas.configure(scrollregion=self.table_canvas.bbox("all")))
        self.table_canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.table_canvas.configure(yscrollcommand=self.scrollbar.set)

        self.table_canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        
        self.table_canvas.bind("<Enter>", self._bind_mousewheel)
        self.table_canvas.bind("<Leave>", self._unbind_mousewheel)

        # 🆕 Updated Headers
        headers = ["Item Number", "Item Name", "Quantity", "Unit", "Last Attachments", "Actions"]
        col_widths = [100, 400, 150, 100, 150, 150]

        header_frame = CTkFrame(self.scrollable_frame, fg_color="black")
        header_frame.grid(row=0, column=0, columnspan=len(headers), sticky="nsew", padx=1, pady=1)

        for col_index, (col_name, col_width) in enumerate(zip(headers, col_widths)):
            col_label = CTkLabel(header_frame, text=col_name, font=("Roboto", 12, "bold"), text_color="white",
                                width=col_width, height=30, fg_color="gray20")
            col_label.grid(row=0, column=col_index, padx=1, pady=1, sticky="nsew")

        for col_index in range(len(headers)):
            self.scrollable_frame.grid_columnconfigure(col_index, weight=1)

        selected_items = self.fetch_selected_construction_items()

        for row_index, item in enumerate(selected_items, start=1):
            bg_color = "#2E2E2E" if row_index % 2 == 0 else "#1C1C1C"
            row_frame = CTkFrame(self.scrollable_frame, fg_color="black")
            row_frame.grid(row=row_index, column=0, columnspan=len(headers), sticky="nsew", padx=1, pady=1)

            # Unpack item info
            item_number, item_name, quantity, unit = item

            # Get last upload date for this item
            last_date = self.get_last_upload_date(item_number, item_name)

            # Fill in data columns
            row_data = [item_number, item_name, quantity, unit, last_date]
            for col_index, (value, col_width) in enumerate(zip(row_data, col_widths[:-1])):
                cell_label = CTkLabel(row_frame, text=value, font=("Roboto", 12), text_color="white",
                                    width=col_width, height=30, fg_color=bg_color)
                cell_label.grid(row=0, column=col_index, padx=1, pady=1, sticky="nsew")

            # Add action button
            add_image_button = CTkButton(
                row_frame, text="Add / Update Image", font=("Roboto", 12),
                command=lambda i=item_number, n=item_name: self.open_add_image_window(i, n),
                width=col_widths[-1], height=30
            )
            add_image_button.grid(row=0, column=len(col_widths) - 1, padx=1, pady=1, sticky="nsew")

    def open_add_image_window(self, item_number, item_name):
        # If an add image window is already open, bring it to the front.
        if self.add_image_window is not None and self.add_image_window.winfo_exists():
            self.add_image_window.lift()
            return

        self.add_image_window = tk.Toplevel()
        self.add_image_window.title(f"Add / Update Image - {item_number} | {item_name}")
        self.add_image_window.geometry("1600x800")
        self.add_image_window.resizable(False, False)
        self.add_image_window.configure(bg="#11151f")
        self.add_image_window.protocol("WM_DELETE_WINDOW", self.on_add_image_window_close)

        # Center the window on the screen.
        self.add_image_window.update_idletasks()
        width, height = 1600, 800
        x = (self.add_image_window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.add_image_window.winfo_screenheight() // 2) - (height // 2)
        self.add_image_window.geometry(f"{width}x{height}+{x}+{y}")

        scrollable_frame = CTkScrollableFrame(self.add_image_window, corner_radius=10, fg_color="#1C1C1C")
        scrollable_frame.pack(fill="both", expand=True, padx=20, pady=(20,0))

        CTkLabel(
            scrollable_frame,
            text="Upload Construction Images",
            font=("Roboto", 20, "bold"),
            text_color="white"
        ).pack(pady=20)

        # --------------------------------------------------
        def has_gps_coordinates(filepath):
            try:
                img = Image.open(filepath)
                exif_data = img._getexif()
                if not exif_data:
                    return False
                exif_tags = {ExifTags.TAGS[k]: v for k, v in exif_data.items() if k in ExifTags.TAGS}
                if "GPSInfo" in exif_tags:
                    gps_info = exif_tags["GPSInfo"]
                    gps_tags = {ExifTags.GPSTAGS.get(k, k): v for k, v in gps_info.items()}
                    return "GPSLatitude" in gps_tags and "GPSLongitude" in gps_tags
                return False
            except Exception as e:
                print(f"Error reading EXIF data: {e}")
                return False

        def set_preview_pic(filepath, label, entry):
            try:
                img = Image.open(filepath)
                img = img.resize((350, 300))
                img = ImageTk.PhotoImage(img)
                label.config(image=img)
                label.image = img
                entry.delete(0, tk.END)
                entry.insert(0, filepath)
            except Exception as e:
                print(f"Error setting preview picture: {e}")

        def select_pic(label, entry):
            filename = filedialog.askopenfilename(
                initialdir=os.getcwd(),
                title="Select Image",
                filetypes=[("Image files", "*.png *.jpg *.jpeg *.jfif")]
            )
            if filename:
                if has_gps_coordinates(filename):
                    set_preview_pic(filename, label, entry)
                else:
                    messagebox.showwarning("Invalid Image", "Please upload a geotagged photo.")

        # --------------------------------------------------
        static_container = CTkFrame(scrollable_frame, corner_radius=10, fg_color="#2B2B2B")
        static_container.pack(pady=10, fill="x", padx=10)

        self.static_section_widgets = {}

        sections = [
            {"title": "Before", "key": "before"},
            {"title": "During", "key": "during"},
            {"title": "After", "key": "after"},
            {"title": "Station Limit/Grid", "key": "station_limits"}
        ]

        for sec in sections:
            card = CTkFrame(static_container, corner_radius=10, fg_color="#3B3B3B")
            card.pack(side="left", expand=True, fill="both", padx=10, pady=10)

            CTkLabel(card, text=sec["title"], font=("Roboto", 16, "bold"), text_color="white").pack(pady=(10, 5))

            if sec["key"] in ["before", "during", "after"]:
                preview_label = tk.Label(card, text="No Image", bg="#3B3B3B", fg="white")
                preview_label.pack(pady=5)
                path_entry = CTkEntry(card, width=200)
                path_entry.pack(pady=5)
                browse_btn = CTkButton(
                    card,
                    text="Browse",
                    width=100,
                    command=lambda lbl=preview_label, ent=path_entry: select_pic(lbl, ent)
                )
                browse_btn.pack(pady=5)
                self.static_section_widgets[sec["key"]] = {"preview_label": preview_label, "entry": path_entry}
            else:
                path_entry = CTkEntry(card, width=200)
                path_entry.pack(pady=(40, 10))
                self.static_section_widgets[sec["key"]] = {"entry": path_entry}

        # Load static data
        try:
            conn = sqlite3.connect(self.db_filename)
            cur = conn.cursor()
            cur.execute("""
                            CREATE TABLE IF NOT EXISTS completed_construction_images (
                                id               INTEGER PRIMARY KEY AUTOINCREMENT,
                                construction_type  TEXT,
                                item_number       TEXT,
                                item_name         TEXT,
                                row_index         INTEGER,
                                image_before      TEXT,
                                image_during      TEXT,
                                image_after       TEXT,
                                station_limits    TEXT,
                                report_generated  INTEGER DEFAULT 0,
                                upload_date       TEXT
                            )
                        """)
            cur.execute(
                "SELECT image_before,image_during,image_after,station_limits FROM completed_construction_images"
                " WHERE item_number=? AND item_name=? AND row_index=0",(item_number,item_name)
            )
            row = cur.fetchone()
            conn.close()
            if row:
                for idx,phase in enumerate(["before","during","after"]):
                    path = row[idx]
                    if path and os.path.exists(path):
                        set_preview_pic(path,
                            self.static_section_widgets[phase]["preview_label"],
                            self.static_section_widgets[phase]["entry"]
                        )
                if row[3]:
                    e = self.static_section_widgets["station_limits"]["entry"]
                    e.delete(0,tk.END); e.insert(0,row[3])
        except Exception as e:
            print(e)

        # --------------------------------------------------
        CTkLabel(scrollable_frame, text="Additional Image Entries",
                 font=("Roboto",16,"bold"),text_color="white").pack(pady=20)
        dynamic_container = CTkFrame(scrollable_frame, corner_radius=10, fg_color="#2B2B2B")
        dynamic_container.pack(pady=10, fill="x", padx=10)
        self.dynamic_rows_entries = []

        def add_new_dynamic_row(prepopulated_data=None):
            row_card = CTkFrame(dynamic_container, corner_radius=10, fg_color="#3B3B3B")
            row_card.pack(fill="x", padx=10, pady=10)
            row_entries = {}

            phases = ["Before", "During", "After", "Station Limit/Grid"]
            for phase in phases:
                phase_frame = CTkFrame(row_card, corner_radius=10, fg_color="#4B4B4B")
                phase_frame.pack(side="left", expand=True, fill="both", padx=10, pady=10)
                CTkLabel(phase_frame, text=phase, font=("Roboto", 12, "bold"), text_color="white").pack(pady=(10, 5))
                if phase in ["Before", "During", "After"]:
                    preview_label = tk.Label(phase_frame, text="No Image", bg="#4B4B4B", fg="white")
                    preview_label.pack(pady=5)
                    path_entry = CTkEntry(phase_frame, width=150)
                    path_entry.pack(pady=5)
                    browse_btn = CTkButton(
                        phase_frame,
                        text="Browse",
                        width=80,
                        command=lambda lbl=preview_label, ent=path_entry: select_pic(lbl, ent)
                    )
                    browse_btn.pack(pady=5)
                    row_entries[phase.lower()] = path_entry
                    if prepopulated_data and prepopulated_data.get(phase.lower(), ""):
                        file_path = prepopulated_data.get(phase.lower())
                        path_entry.insert(0, file_path)
                        if os.path.exists(file_path):
                            set_preview_pic(file_path, preview_label, path_entry)
                else:
                    path_entry = CTkEntry(phase_frame, width=150)
                    path_entry.pack(pady=(40, 10))
                    row_entries["station_limits"] = path_entry
                    if prepopulated_data and prepopulated_data.get("station_limits", ""):
                        path_entry.insert(0, prepopulated_data.get("station_limits"))
            remove_btn = CTkButton(
                row_card,
                text="Remove",
                width=50,
                fg_color="red",
                hover_color="darkred",
                command=lambda: remove_dynamic_row(row_card, row_entries)
            )
            remove_btn.pack(side="left", padx=10, pady=10)
            self.dynamic_rows_entries.append(row_entries)

        def remove_dynamic_row(row_card, row_entries):
            row_card.destroy()
            if row_entries in self.dynamic_rows_entries:
                self.dynamic_rows_entries.remove(row_entries)

        # load saved construction rows
        try:
            conn = sqlite3.connect(self.db_filename)
            cur = conn.cursor()
            cur.execute("SELECT image_before,image_during,image_after,station_limits"
                        " FROM completed_construction_images"
                        " WHERE item_number=? AND item_name=? AND row_index>0",
                        (item_number,item_name))
            for data in cur.fetchall():
                pre = {"before":data[0],"during":data[1],"after":data[2],"station_limits":data[3]}
                add_new_dynamic_row(pre)
            conn.close()
        except Exception:
            pass

        # Initialize testing rows list
        self.testing_rows = []

        # --------------------------------------------------
        # Helpers for testing rows

        def select_testing_images(testing_files_list, preview_button):
            files = filedialog.askopenfilenames(
                initialdir=os.getcwd(),
                title="Select Testing Images",
                filetypes=[("Image files", "*.png *.jpg *.jpeg *.jfif")]
            )
            if files:
                for f in files:
                    if f not in testing_files_list:
                        testing_files_list.append(f)
                preview_button.configure(state="normal")

        def remove_testing_row(row_card, testing_dict):
            row_card.destroy()
            if testing_dict in self.testing_rows:
                self.testing_rows.remove(testing_dict)
        
        def preview_testing_images(files):
            if not files: return
            win = tk.Toplevel(self.add_image_window)
            win.title("Preview Testing Images")
            win.geometry("800x600")

            canvas = tk.Canvas(win, bg="#1C1C1C")
            vsb    = tk.Scrollbar(win, orient="vertical", command=canvas.yview)
            canvas.configure(yscrollcommand=vsb.set)
            vsb.pack(side="right", fill="y")
            canvas.pack(side="left", fill="both", expand=True)

            frame = tk.Frame(canvas, bg="#1C1C1C")
            canvas.create_window((0,0), window=frame, anchor="nw")
            frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

            for idx, path in enumerate(files):
                try:
                    img = Image.open(path)
                    img.thumbnail((200, 150))
                    photo = ImageTk.PhotoImage(img)
                    lbl = tk.Label(frame, image=photo, bg="#1C1C1C")
                    lbl.image = photo
                    lbl.grid(row=idx//3, column=idx%3, padx=5, pady=5)
                except Exception as e:
                    print(f"Cannot preview {path}: {e}")

        def add_testing_image_row(prepopulated_name=None, prepopulated_files=None):
            row_card = CTkFrame(dynamic_container, corner_radius=10, fg_color="#3B3B3B")
            row_card.pack(fill="x", padx=10, pady=10)

            # Name entry
            name_entry = CTkEntry(row_card, width=200, placeholder_text="Testing Name")
            name_entry.pack(side="left", padx=(10,5), pady=10)
            if prepopulated_name:
                name_entry.insert(0, prepopulated_name)

            # File list & preview button
            files = list(prepopulated_files) if prepopulated_files else []
            preview_btn = CTkButton(
                row_card, text="Preview Images", width=120,
                state="normal" if files else "disabled",
                command=lambda: preview_testing_images(files)
            )
            preview_btn.pack(side="left", padx=5, pady=10)

            # Browse
            CTkButton(
                row_card, text="Browse Images", width=100,
                command=lambda: select_testing_images(files, preview_btn)
            ).pack(side="left", padx=5, pady=10)

            # Remove
            remove_btn = CTkButton(
                row_card, text="Remove", width=60, fg_color="red", hover_color="darkred",
                command=lambda: remove_testing_row(row_card, testing_dict)
            )
            remove_btn.pack(side="left", padx=5, pady=10)

            testing_dict = {"name_entry": name_entry, "files": files}
            self.testing_rows.append(testing_dict)

        # --------------------------------------------------
        
        def update_all_uploaded_images(item_number, item_name, static_data, dynamic_rows):
            base_dir = os.path.join(os.getcwd(), "images")
            conn = sqlite3.connect(self.db_filename)
            cursor = conn.cursor()

            # --- ensure tables exist ---
            cursor.execute("""
            CREATE TABLE IF NOT EXISTS completed_construction_images (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                construction_type TEXT,
                item_number TEXT,
                item_name TEXT,
                row_index INTEGER,
                image_before TEXT,
                image_during TEXT,
                image_after TEXT,
                station_limits TEXT,
                report_generated INTEGER DEFAULT 0,
                upload_date TEXT
            )""")
            cursor.execute("""
            CREATE TABLE IF NOT EXISTS testing_images (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                item_number TEXT,
                item_name TEXT,
                test_index INTEGER,
                test_name TEXT,
                image_path TEXT,
                upload_date TEXT
            )""")

            # --- pull existing rows (so we can compare and delete replaced files) ---
            cursor.execute("""
            SELECT row_index, image_before, image_during, image_after, station_limits, report_generated
            FROM completed_construction_images
            WHERE item_number=? AND item_name=?
            """, (item_number, item_name))
            existing = {
                r[0]: (r[1], r[2], r[3], r[4], r[5])
                for r in cursor.fetchall()
            }
            # --- pull existing testing-image paths ---
            cursor.execute("""
            SELECT image_path
            FROM testing_images
            WHERE item_number=? AND item_name=?
            """, (item_number, item_name))
            old_testing_paths = [r[0] for r in cursor.fetchall()]

            # --- wipe old records ---
            cursor.execute("DELETE FROM completed_construction_images WHERE item_number=? AND item_name=?",
                        (item_number, item_name))
            cursor.execute("DELETE FROM testing_images WHERE item_number=? AND item_name=?",
                        (item_number, item_name))

            today = datetime.date.today().isoformat()

            def copy_new_file(src, phase):
                """Copy src into base_dir/phase, with a uuid suffix."""
                tgt_dir = os.path.join(base_dir, phase)
                os.makedirs(tgt_dir, exist_ok=True)
                name = os.path.basename(src)
                uniq = f"{os.path.splitext(name)[0]}_{uuid.uuid4().hex}{os.path.splitext(name)[1]}"
                dst = os.path.join(tgt_dir, uniq)
                shutil.copy(src, dst)
                return dst

            def handle_image(src, old_path, phase):
                """
                If src is already in our images folder, reuse; 
                else copy new and delete old_path if it existed.
                """
                if not src:
                    return None

                abs_src = os.path.abspath(src)
                abs_base = os.path.abspath(base_dir)

                # 1) reused
                if abs_src.startswith(abs_base) and os.path.exists(abs_src):
                    return abs_src

                # 2) brand new
                new_path = copy_new_file(src, phase)

                # 3) clean up old file if it's different
                if old_path and old_path != new_path and os.path.exists(old_path):
                    os.remove(old_path)

                return new_path

            # --- static row (row_index=0) ---
            old0 = existing.get(0, (None, None, None, None, 0))
            new0 = {}
            for i, phase in enumerate(("before", "during", "after")):
                new0[phase] = handle_image(
                    static_data.get(phase),
                    old0[i],
                    phase
                )
            new0["station_limits"] = static_data.get("station_limits", "")
            
            # … after computing new0 …
            report_flag = old0[4] if (
                old0[:4] == (
                    new0["before"],
                    new0["during"],
                    new0["after"],
                    new0["station_limits"]
                )
            ) else 0

            # --- insert static row ---
            cursor.execute("""
            INSERT INTO completed_construction_images
                (item_number, item_name, row_index,
                image_before, image_during, image_after, station_limits,
                report_generated, upload_date)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
            item_number, 
            item_name, 
            0,
            new0["before"], 
            new0["during"], 
            new0["after"],
            new0["station_limits"],
            report_flag,
            today
            ))

            # --- dynamic rows (row_index 1..) ---
            for idx, row in enumerate(dynamic_rows, start=1):
                oldr = existing.get(idx, (None, None, None, None, 0))
                newr = {}
                for i, phase in enumerate(("before", "during", "after")):
                    val = row[phase].get().strip()
                    newr[phase] = handle_image(val, oldr[i], phase)
                newr["station_limits"] = row["station_limits"].get().strip()
                flag = oldr[4] if oldr[:4] == (
                    newr["before"],
                    newr["during"],
                    newr["after"],
                    newr["station_limits"]
                ) else 0

                cursor.execute("""
                INSERT INTO completed_construction_images
                    (item_number, item_name, row_index,
                    image_before, image_during, image_after, station_limits,
                    report_generated, upload_date)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                item_number,
                item_name,
                idx,
                newr["before"],
                newr["during"],
                newr["after"],
                newr["station_limits"],
                flag,
                today
                ))

            # --- testing rows ---
            new_testing_paths = []
            for t_idx, testing in enumerate(self.testing_rows, start=1):
                name = testing["name_entry"].get().strip()
                for fp in testing["files"]:
                    tp = handle_image(fp, None, "testing")
                    new_testing_paths.append(tp)
                    cursor.execute("""
                    INSERT INTO testing_images
                        (item_number, item_name, test_index, test_name, image_path, upload_date)
                    VALUES (?, ?, ?, ?, ?, ?)
                    """, (item_number, item_name, t_idx, name, tp, today))

            # --- delete any old testing-image files the user didn’t re-upload ---
            for old_fp in old_testing_paths:
                if old_fp not in new_testing_paths and os.path.exists(old_fp):
                    os.remove(old_fp)

            conn.commit()
            conn.close()

        def upload_images():
            static_data = {
                "before": self.static_section_widgets["before"]["entry"].get().strip(),
                "during": self.static_section_widgets["during"]["entry"].get().strip(),
                "after": self.static_section_widgets["after"]["entry"].get().strip(),
                "station_limits": self.static_section_widgets["station_limits"]["entry"].get().strip()
            }
            static_provided = static_data["before"] or static_data["during"] or static_data["after"]
            dynamic_provided = any(
                row["before"].get().strip() or row["during"].get().strip() or row["after"].get().strip()
                for row in self.dynamic_rows_entries
            )
            if not (static_provided or dynamic_provided):
                messagebox.showerror("Error", "Please upload at least one geotagged image.")
                return
            try:
                update_all_uploaded_images(
                    item_number, item_name,
                    static_data=static_data,
                    dynamic_rows=self.dynamic_rows_entries
                )
                messagebox.showinfo("Success", "Images and station limits updated successfully!")
                self.refresh_table()
                self.on_add_image_window_close()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to update images: {e}")
                
        # --- Load saved testing rows ---
        try:
            conn = sqlite3.connect(self.db_filename)
            cur = conn.cursor()
            cur.execute(
                "SELECT test_index,test_name,image_path FROM testing_images"
                " WHERE item_number=? AND item_name=? ORDER BY test_index,id",
                (item_number, item_name)
            )
            grouped = defaultdict(lambda: {"name":None, "files":[]})
            for idx,name,path in cur.fetchall():
                grouped[idx]["name"] = name
                grouped[idx]["files"].append(path)
            conn.close()
            for idx in sorted(grouped):
                add_testing_image_row(
                    prepopulated_name=grouped[idx]["name"],
                    prepopulated_files=grouped[idx]["files"]
                )
        except Exception as e:
            print(e)

        # --- Bottom controls ---
        bottom_frame = CTkFrame(self.add_image_window, corner_radius=10, fg_color="#2B2B2B")
        bottom_frame.pack(fill="x", padx=20, pady=20)
        
        action_frame = CTkFrame(bottom_frame, corner_radius=10, fg_color="#2B2B2B")
        action_frame.pack(side="left")
        CTkButton(action_frame, text="Add Row", width=150,
                command=add_new_dynamic_row).pack(side="left", padx=5, pady=5)
        CTkButton(action_frame, text="Add Image for Testing", width=150,
                command=add_testing_image_row).pack(side="left", padx=5, pady=5)

        btn_frame = CTkFrame(bottom_frame, corner_radius=10, fg_color="#2B2B2B")
        btn_frame.pack(side="right")
        CTkButton(btn_frame, text="Submit", width=150,
                command=upload_images, fg_color="#239409", hover_color="#1e7f0d")\
            .pack(side="left", padx=5)
        CTkButton(btn_frame, text="Cancel", width=150,
                fg_color="#b00505", hover_color="#8f0404",
                command=self.on_add_image_window_close)\
            .pack(side="left", padx=5)
        CTkButton(btn_frame, text="Generate Report", width=150,
                command=lambda: self.generate_report(item_number, item_name))\
            .pack(side="left", padx=5)

    def on_add_image_window_close(self):
        if self.add_image_window is not None:
            self.add_image_window.destroy()
            self.add_image_window = None

    def close_add_image_window(self):
        if self.add_image_window is not None and self.add_image_window.winfo_exists():
            self.add_image_window.destroy()
            self.add_image_window = None

    def fetch_project_details(self):
        try:
            conn = sqlite3.connect(self.db_filename)
            cursor = conn.cursor()
            cursor.execute("SELECT project_name, location, project_id, contractor_name, project_type FROM project_informations LIMIT 1")
            row = cursor.fetchone()
            return dict(zip(["project_name", "location", "project_id", "contractor_name", "project_type"], row)) if row else {}
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error fetching project details: {e}")
            return {}
        finally:
            conn.close()

    def fetch_selected_construction_items(self):
        try:
            conn = sqlite3.connect(self.db_filename)
            cursor = conn.cursor()
            # Modified query to fetch item_number, item_name, quantity, unit
            cursor.execute("""
                SELECT item_number, item_name, quantity, unit
                FROM selected_construction_items
            """)
            return cursor.fetchall()
        except sqlite3.Error as e:
            messagebox.showerror("Database Error", f"Error fetching construction items: {e}")
            return []
        finally:
            if conn:
                conn.close()

    def generate_report(self, item_number, item_name):
        """
        Generate a single Word document report combining construction images (before/during/after/station limits)
        and materials testing images for a given item.
        """
        # Connect and fetch data
        try:
            conn = sqlite3.connect(self.db_filename)
            cursor = conn.cursor()
            # Static row
            cursor.execute(
                "SELECT image_before, image_during, image_after, station_limits, report_generated"
                " FROM completed_construction_images"
                " WHERE item_number=? AND item_name=? AND row_index=0",
                (item_number, item_name)
            )
            static_row = cursor.fetchone()
            # Dynamic rows
            cursor.execute(
                "SELECT row_index, image_before, image_during, image_after, station_limits, report_generated"
                " FROM completed_construction_images"
                " WHERE item_number=? AND item_name=? AND row_index>0",
                (item_number, item_name)
            )
            dynamic_rows = cursor.fetchall()
            # Testing rows
            cursor.execute(
                "SELECT test_name, image_path FROM testing_images"
                " WHERE item_number=? AND item_name=?"
                " ORDER BY test_index, id",
                (item_number, item_name)
            )
            testing_rows = cursor.fetchall()
            # Project info
            cursor.execute(
                "SELECT project_id, project_name, location, contractor_name"
                " FROM project_informations LIMIT 1"
            )
            proj = cursor.fetchone()
            conn.close()
            if not proj:
                messagebox.showerror("Error", "Project information not found.")
                return
            project_id, project_name, location, contractor_name = proj
        except Exception as e:
            messagebox.showerror("Error", f"Data fetch failed: {e}")
            return

        # Build document
        document = Document()
        document.styles['Normal'].paragraph_format.space_after = Pt(0)
        for sec in document.sections:
            sec.top_margin = Cm(1.27)
            sec.bottom_margin = Cm(1.27)
            sec.left_margin = Cm(1.27)
            sec.right_margin = Cm(1.27)
        # Header (logos/text)
        section = document.sections[0]
        header = section.header
        width = section.page_width - section.left_margin - section.right_margin
        tbl = header.add_table(rows=1, cols=3, width=width)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c0, c1, c2 = tbl.rows[0].cells
        set_cell_width(c0, 2000); set_cell_width(c1, 8000); set_cell_width(c2, 2000)
        # Left logo
        left_logo = resource_path("images/prdp_logo.png")
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if os.path.exists(left_logo): p0.add_run().add_picture(left_logo, width=Inches(1))
        else: p0.add_run("Left Logo")
        # Center text
        p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run("Republic of the Philippines\n"); run = p1.add_run("PHILIPPINE RURAL DEVELOPMENT PROJECT\n"); run.bold=True
        p1.add_run("(Input Province Name)\n"); run2=p1.add_run("(Input Municipality Name)"); run2.bold=True
        # Right logo placeholder
        nested = c2.add_table(rows=1, cols=1); nc = nested.cell(0,0)
        p2 = nc.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run("Insert Municipality Logo")

        # Combine rows
        rows = []
        if static_row: rows.append((0, *static_row))
        rows.extend(dynamic_rows)
        # Add construction images section
        for idx, row in enumerate(rows):
            if idx>0: document.add_page_break()
            # Project header
            for label, value in [
                ("NAME OF PROJECT:", project_name),
                ("LOCATION:", location),
                ("PROJECT ID:", project_id),
                ("CONTRACTOR:", contractor_name)
            ]:
                p = document.add_paragraph(); p.add_run(f"{label} ").bold=True; p.add_run(str(value))
            document.add_paragraph()
            # Item title
            p = document.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"ITEM {item_number.upper()} {item_name.upper()}").bold=True
            # Static vs dynamic unpack
            if row[0]==0:
                _, before_img, during_img, after_img, station_limits, _ = row
            else:
                _, before_img, during_img, after_img, station_limits, _ = row
            if station_limits:
                ps = document.add_paragraph(f"STATION LIMIT: {station_limits}"); ps.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            # Phases
            for phase, img in zip(["BEFORE","DURING","AFTER"],[before_img,during_img,after_img]):
                pph = document.add_paragraph(phase); pph.alignment=WD_ALIGN_PARAGRAPH.CENTER
                if img and os.path.exists(img):
                    try:
                        document.add_picture(img, width=Inches(4), height=Inches(2)); last=document.paragraphs[-1]; last.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        document.add_paragraph("Error adding image.")
                else:
                    document.add_paragraph("No image available.")

        # Add testing section if present
        if testing_rows:
            document.add_page_break()
            from collections import defaultdict
            grp = defaultdict(list)
            for name, path in testing_rows: grp[name].append(path)
            for t_idx, (test_name, paths) in enumerate(grp.items()):
                if t_idx>0: document.add_page_break()
                # Project header
                for label, value in [
                    ("NAME OF PROJECT:", project_name),
                    ("LOCATION:", location),
                    ("PROJECT ID:", project_id),
                    ("CONTRACTOR:", contractor_name)
                ]:
                    p = document.add_paragraph(); p.add_run(f"{label} ").bold=True; p.add_run(str(value))
                document.add_paragraph()
                # Test title
                pt = document.add_paragraph(); pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
                pt.add_run(f"{test_name.upper()}").bold=True
                # Images
                for img in paths:
                    if img and os.path.exists(img):
                        try:
                            document.add_picture(img, width=Inches(4), height=Inches(3)); lp=document.paragraphs[-1]; lp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                        except:
                            document.add_paragraph("Error adding image.")
                    else:
                        document.add_paragraph("No image available.")

        # Save file
        save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                        filetypes=[("Word Documents","*.docx")], title="Save Report As")
        if save_path:
            try:
                document.save(save_path)
                messagebox.showinfo("Success", f"Report saved as {save_path}")
                # mark construction rows reported
                conn = sqlite3.connect(self.db_filename); cur = conn.cursor()
                if static_row and static_row[4]==0:
                    cur.execute("UPDATE completed_construction_images SET report_generated=1 WHERE item_number=? AND item_name=? AND row_index=0",
                                (item_number,item_name))
                for dr in dynamic_rows:
                    if dr[5]==0:
                        cur.execute("UPDATE completed_construction_images SET report_generated=1 WHERE item_number=? AND item_name=? AND row_index=?",
                                    (item_number,item_name,dr[0]))
                conn.commit(); conn.close()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save document: {e}")

    def on_close(self):
        self.logout()

    def logout(self):
        self.close_add_image_window()
        self.logout_button.configure(state="disabled")
        if self.login_app:
            self.login_app.deiconify()
        self.withdraw()
        self.after(200, self.destroy)
